import csv, os, webbrowser
import docx
# fund_fees is a dictionary with the key
# being the health fund ID and the data being a list with the consult fee
# and the unit fee for each fund.
# also import medicare codes
from module.fundfees import fund_fees, pe_code, col_code, age_code, sick_code

# these html codes for the account
# from module.templates import base_html, table_html, terminal_html

while True:
    try:
        csvfile = open('/Users/jtair/Downloads/JT Patients 2016 (Responses) - Form responses 1.csv', 'r')
    except IOError:
        url = 'https://docs.google.com/spreadsheets/d/13p-ZYuUlmxo9fqnxru1Pg06xLES0gt_TwPJAw5nZrYU/edit?usp=sharing'
        webbrowser.open(url)
        break

    patlist = csv.reader(csvfile)
    number_to_print = 0         # output the number of accounts printed as a check
    input_day = input("input day in format --   ")
    input_month = input("input month in format --   ")
    doc = docx.Document()
    for row in patlist:
        if input_day == row[0][0:2] and input_month == row[0][3:5]\
                and row[2] not in ['BB', 'GARRISON', 'OS', 'VA']:
            # extract data from csv file into variables - these are strings
            ep_date = row[0][0:2] + '-' + row[0][3:5] + '-' + row[0][6:10]
            patient = row[1]
            fund = row[2]
            endo = row[3]
            colon = row[4]
            age = row[5]
            sick = row[6]
            time = row[7]
            doctor = row[8]

            number_to_print += 1

            # get fees from module depending on fund
            fee_package = fund_fees[fund]
            consult = fee_package[0]
            consult_as_float = float(fee_package[0])
            unit = float(fee_package[1])

            # get time info and calculate time fee
            # the fourth digit in the time code gives the number of units
            time_length = int(time[3])
            time_fee = time_length * unit

            # calculate total_fee, initialise total_fee and add on consult
            total_fee = consult_as_float

            if endo == 'Yes':
                total_fee += (unit * 5)
            if endo == 'No' and colon == 'Yes':
                total_fee += (unit * 4)
            if age == 'Yes':
                total_fee += unit
            if sick == 'Yes':
                total_fee += unit

            # add on time fees
            total_fee = total_fee + (time_length * unit)


            doc.add_heading('Account for Anaesthetic',level = 0)
            doc.add_heading('Dr John Tillett',level = 2)
            doc.add_heading('7 Henry Lawson Drive, Villawood NSW 2163',level = 4)
            doc.add_heading('Provider Number: 0307195H',level = 5)
            doc.add_heading('Phone: 8382 6622  Email: john@endoscopy.stvincents.com.au', level = 5)
            doc.add_paragraph('')
            # p_prov = doc.add_paragraph('Provider Number: ')
            # p_prov.add_run('0307195H').bold = True

            # p_enq = doc.add_paragraph('Enquiries:')
            # p_enq.add_run('  phone').italic = True
            # p_enq.add_run('  0408 116 320')
            # p_enq.add_run('  fax').italic = True
            # p_enq.add_run('  02 8382 6622')
            # p_enq.add_run('  email').italic = True
            # p_enq.add_run('  tillett1957@gmail.com')
            # doc.add_paragraph('Patient Details')
            doc.add_paragraph('')
            doc.add_paragraph('%s' % patient)
            doc.add_paragraph('')
            doc.add_paragraph('%s' % fund)
            doc.add_paragraph('')
            doc.add_paragraph('Date of Procedure:  %s' % ep_date)
            doc.add_paragraph('Procedure performed by Dr %s at the Diagnostic Endoscopy Centre, Darlinghurst, NSW 2010' % doctor)
            doc.add_paragraph('')

            doc.add_paragraph('Item Number%sFee' % (' ' * 10))

            p_cons = doc.add_paragraph('17610')
            cons_str = '%.2f' % consult_as_float
            cons_str = cons_str.rjust(25)
            p_cons.add_run(cons_str)

            if endo == 'Yes':
                p_endo = doc.add_paragraph(pe_code)
                endo_str = '%.2f' % (unit * 5)
                endo_str = endo_str.rjust(24)
                p_endo.add_run(endo_str)
                if colon == 'Yes':
                    p_col = doc.add_paragraph(col_code)
                    col_str = '%.2f' % 0.0
                    col_str = col_str.rjust(26)
                    p_col.add_run(col_str)
            if endo == 'No' and colon == 'Yes':
                p_col = doc.add_paragraph(col_code)
                col_str = '%.2f' % (unit * 4)
                col_str = col_str.rjust(24)
                p_col.add_run(col_str)
            if age == 'Yes':
                p_age = doc.add_paragraph(age_code)
                age_str = '%.2f' % unit
                age_str = age_str.rjust(25)
                p_age.add_run(age_str)
            if sick == 'Yes':
                p_sick = doc.add_paragraph(sick_code)
                sick_str = '%.2f' % unit
                sick_str = sick_str.rjust(25)
                p_sick.add_run(sick_str)

            p_time_fee = doc.add_paragraph(time)
            time_fee_str = '%.2f' % time_fee
            time_fee_str = time_fee_str.rjust(25)
            p_time_fee.add_run(time_fee_str)

            p_tot = doc.add_paragraph('Total Fee')
            tot_str = '$%.2f' % total_fee
            tot_str = tot_str.rjust(19)
            p_tot.add_run(tot_str)

            doc.add_paragraph('')
            p_gst = doc.add_paragraph('')
            p_gst.add_run('No item on this invoice attracts GST').italic=True

            doc.add_page_break()


    doc.save('/Users/jtair/Documents/Invoices/accts.docx')
    csvfile.close()
    # final print out of number of accounts - done as a check

    if number_to_print == 0:
        print('***There seems to be no patients that day.***')
        print('     ***********')
    else:
        print('Number of accounts printed is {}'.format(number_to_print))
        print('      **********')
    # more = input('Print more accounts? y or n :')
    # if more == 'n':
    #     os.remove('/Users/jtair/Downloads/JT Patients 2016 (Responses) - Form responses 1.csv')
    #     break
